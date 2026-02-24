import fitz  # PyMuPDF
import os
from docx import Document
from docx.shared import Pt
from PIL import Image
import io

# --- CONFIGURAÇÕES ---
pdf_file = "entrada.pdf"
docx_file = "saida_pymupdf.docx"

# --- Função para limpar caracteres inválidos ---
def clean_text(text: str) -> str:
    """Remove caracteres inválidos para XML/Word (ex: NULL bytes e controles)."""
    if not text:
        return ""
    # Mantém apenas caracteres válidos para XML
    return "".join(
        ch for ch in text
        if ch == "\t" or ch == "\n" or ch == "\r" or ord(ch) >= 32
    )

def convert_with_pymupdf(pdf_path, docx_path):
    """
    Converte PDF para DOCX usando PyMuPDF para extrair conteúdo e python-docx
    para construir o documento.
    """
    if not os.path.exists(pdf_path):
        print(f"Erro: O arquivo de entrada '{pdf_path}' não foi encontrado.")
        return

    print(f"Iniciando a conversão de '{pdf_path}' com PyMuPDF...")
    doc = Document()
    pdf_document = fitz.open(pdf_path)

    print(f"O PDF tem {len(pdf_document)} páginas.")

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        print(f"Processando página {page_num + 1}...")

        # Extrai todos os elementos da página (texto e imagens) em blocos
        blocks = page.get_text("dict")["blocks"]

        # Extrai imagens
        image_list = page.get_images(full=True)

        # Processa blocos de texto
        for block in blocks:
            if block['type'] == 0:  # 0 = bloco de texto
                for line in block["lines"]:
                    if not line["spans"]:
                        continue
                    p = doc.add_paragraph()
                    font_size = line["spans"][0].get("size", 11)
                    for span in line["spans"]:
                        safe_text = clean_text(span.get("text", ""))
                        if not safe_text.strip():
                            continue
                        run = p.add_run(safe_text)
                        run.font.size = Pt(font_size)
                        # Mantém negrito/itálico quando possível
                        if "bold" in span["font"].lower():
                            run.bold = True
                        if "italic" in span["font"].lower():
                            run.italic = True

        # Processa imagens
        if image_list:
            print(f"Encontradas {len(image_list)} imagens na página {page_num + 1}.")
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]

                try:
                    image = Image.open(io.BytesIO(image_bytes))
                    doc.add_picture(io.BytesIO(image_bytes))
                except Exception as e:
                    print(f"  - Não foi possível processar a imagem {img_index}: {e}")

        # Adiciona quebra de página, exceto para a última
        if page_num < len(pdf_document) - 1:
            doc.add_page_break()

    try:
        doc.save(docx_path)
        print(f"✅ Conversão concluída! Arquivo salvo como: '{docx_path}'")
    except Exception as e:
        print(f"❌ Erro ao salvar o arquivo DOCX: {e}")

if __name__ == "__main__":
    convert_with_pymupdf(pdf_file, docx_file)
